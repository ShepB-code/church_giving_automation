from docx import Document
from docx.shared import Pt
import json

from dotenv import load_dotenv

load_dotenv()

import os

def write_doc(doc_name):
    all_info = read_json()
    document = Document()
    for donor in all_info.keys():
       generate_page(donor, all_info[donor], document, "2020") 

    document.save(f"{doc_name}.docx")
def generate_page(name, donor_info, word_document, year):
    total = "{:.2f}".format(donor_info["Total"])
    text = f"\n\n\n\n\n\nDear {name}, \n\tOn behalf of {os.environ.get('CHURCH_NAME')}, we would like to thank you for the support you have given this past year. Your total tithes and gifts for {year} were ${total}. Below is a detailed history of your giving.\n\n\n\n"

    #word_document.add_paragraph(text)
    paragraph = word_document.add_paragraph().add_run(text)
    font = paragraph.font
    font.size = Pt(14)
    
    table = word_document.add_table(rows=0, cols=3)
    table.style = 'LightShading-Accent1'

    column_headers = table.add_row().cells
    for table_header, header_label in zip(column_headers, ["Date", "Check/Type", "Amount"]): #Add dollar sign
        table_header.text = header_label 

    general_giving_total = 0
    benevolence_giving_total = 0


    max_len = get_max_len(donor_info)
    #Left justify the numbers!
    #Add dollar sign
    for key, value in donor_info.items():
        if key == 'Total':
            continue
        for date, info in value.items():
            
            if key == "Giving":
                general_giving_total += info[1]
            elif key == "Benevolence":
                benevolence_giving_total += info[1]
            cells = table.add_row().cells
            cells[0].text = date
            # cells[1].text = info[0]
            # cells[2].text = info[1]
            for pog in zip(info, cells[1:]):
                
                if "#" not in str(pog[0]) and "Cash" not in str(pog[0]) and "cash" not in str(pog[0]):
                    amount = str(pog[0])
                    amount = "$" + ("  " * (max_len - len(amount))) + "{:.2f}".format(pog[0])
                    pog[1].text = amount
                else:
                    pog[1].text = str(pog[0])
            # check_num = info[0]
            # amount = info[1]\
    
    #Rounding
    general_giving_total = "{:.2f}".format(general_giving_total)
    benevolence_giving_total = "{:.2f}".format(benevolence_giving_total)
    summary_text = f"\n\n\n\nSummary:\n\tGeneral Giving: ${general_giving_total}\n\tBenevolence Fund: ${benevolence_giving_total}\n\tTotal Giving: ${total}"

    paragraph = word_document.add_paragraph().add_run(summary_text)
    font = paragraph.font
    font.size = Pt(14)

    
    
    word_document.add_page_break()

def get_max_len(dict1):
    max_len = 0
    for key, value in dict1.items():
        if key == "Total":
            continue
        for date, info in value.items():
            if len(str(info[1])) > max_len:
                max_len = len(str(info[1]))

    return max_len
def read_json():
    with open("giving.json", "r") as f:
        return json.load(f)

